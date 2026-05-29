"""
python-docx를 사용한 Word 문서 생성.
"""

import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def to_docx(mode: str, data: dict) -> io.BytesIO:
    if mode == "meeting":
        return _meeting_docx(data)
    return _field_docx(data)


def _heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _bullet(doc: Document, text: str):
    doc.add_paragraph(text, style="List Bullet")


def _meeting_docx(data: dict) -> io.BytesIO:
    doc = Document()
    doc.add_heading("회의록", 0)
    doc.add_paragraph(f"작성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph()

    _heading(doc, "1. 안건", 1)
    for item in data.get("agenda", []):
        _bullet(doc, item)

    _heading(doc, "2. 결정사항", 1)
    for item in data.get("decisions", []):
        _bullet(doc, item)

    _heading(doc, "3. 액션아이템", 1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "업무 내용"
    hdr[1].text = "담당자"
    hdr[2].text = "기한"
    for ai in data.get("action_items", []):
        row = table.add_row().cells
        row[0].text = ai.get("task", "")
        row[1].text = ai.get("assignee") or "-"
        row[2].text = ai.get("due_date") or "-"

    _heading(doc, "4. 논의 요약", 1)
    doc.add_paragraph(data.get("summary", ""))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _field_docx(data: dict) -> io.BytesIO:
    doc = Document()
    doc.add_heading("현장 메모", 0)

    tags = data.get("tags", {})
    meta_lines = []
    if tags.get("project_name"):
        meta_lines.append(f"프로젝트: {tags['project_name']}")
    if tags.get("location"):
        meta_lines.append(f"위치: {tags['location']}")
    if tags.get("record_date"):
        meta_lines.append(f"날짜: {tags['record_date']}")
    meta_lines.append(f"작성일시: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("\n".join(meta_lines))
    doc.add_paragraph()

    _heading(doc, "1. 업무 지시사항", 1)
    for item in data.get("instructions", []):
        _bullet(doc, item)

    _heading(doc, "2. 확인 필요 항목", 1)
    for item in data.get("check_items", []):
        _bullet(doc, item)

    _heading(doc, "3. 현장 관찰 메모", 1)
    for item in data.get("observations", []):
        _bullet(doc, item)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
